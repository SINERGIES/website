from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("outputs/inovia_cr")
DOCX_OUT = OUT_DIR / "CR_INOVIA_hebdomadaires_2026.docx"
MD_OUT = OUT_DIR / "CR_INOVIA_hebdomadaires_2026.md"
INDIVIDUAL_DIR = OUT_DIR / "individuels"

BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
BORDER = "808080"


MEETINGS = [
    {
        "date": "05.01.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Organisation des réunions",
                "discussion": [
                    "Plusieurs indisponibilités et jours fériés affectent les réunions du lundi : absence de PF les 19/01 et 30/03, lundi de Pâques le 06/04, Pentecôte le 25/05.",
                    "Un lien de visioconférence doit être préparé pour les prochaines réunions.",
                ],
                "followup": "Créer le lien de visioconférence.",
            },
            {
                "subject": "Financements et appels à projets",
                "discussion": [
                    "La Fondation Béatrice Denys est écartée.",
                    "L’appel à projets RAISE pourrait financer une étude rétrospective sur deux EDS de la région, avec un budget annoncé de 60 kEUR.",
                    "L’axe proposé porte sur un dictionnaire d’harmonisation des variables obstétriques pour des travaux d’IA : choix des variables, qualification des cas et construction d’un métamodèle de prédiction.",
                    "L’axe données de santé pour RAISE consisterait à récupérer des situations de naissance avec pH < 7, enregistrement CTG et contexte obstétrical, notamment sur cinq ans de données d’enfants transférés.",
                ],
                "followup": "Demander à Ingrid Tissot l’état de l’EDS et son éligibilité pour RAISE.",
            },
            {
                "subject": "Variables et issues néonatales",
                "discussion": [
                    "Les variables d’intérêt clinique a priori doivent être identifiées, ainsi que les variables à exclure en raison de leur faible présence.",
                    "Les issues potentielles discutées sont : acidose/pH, variabilité du FHR, score d’Apgar, excès ou déficit de base, transfert du nouveau-né en soins intensifs, équilibre acido-basique, lactates et température néonatale.",
                    "Une issue défavorable pourrait être définie par transfert, pH < 7 et/ou Apgar à 5 minutes < 7.",
                    "Un score clinico-biologique combinant Apgar et équilibre acido-basique est envisagé, avec une pondération prioritairement clinique.",
                    "Le FHR serait utilisé comme outil de prédiction des paramètres postnataux ; les paramètres envisagés sont la ligne de base, la variabilité et le rythme moyen.",
                ],
                "followup": "Vérifier avec Odile la description des variables et les champs parto S1/S2.",
            },
            {
                "subject": "Consortium INTERREG",
                "discussion": [
                    "La composition du consortium reste à clarifier : CHUB, UMLP, CHUV, UNIGE et entreprises suisses potentielles.",
                    "Deux options sont évoquées : INOVIA + universités, ou deux grands groupes + universités + INOVIA.",
                    "Les contacts HUG d’OB et les relais industriels possibles restent à préciser. NM connaît notamment Samsung via Benjamin Tarterre.",
                ],
                "followup": "Identifier les partenaires suisses et demander une réunion de cadrage INTERREG.",
            },
            {
                "subject": "Article et base de données",
                "discussion": [
                    "L’objectif scientifique cité est de tester si l’ajout de variables cliniques améliore un système de CTG informatisé pour prédire une acidémie néonatale sévère.",
                    "DeepCTG est décrit comme une régression logistique et non comme un modèle de deep learning.",
                    "Après validation de la base synthétique, deux suites sont envisagées : passer par un PHRC pour tester sur d’autres centres et ajouter des cas pathologiques via ces centres.",
                ],
                "followup": "",
            },
        ],
    },
    {
        "date": "12.01.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Administration et propriété intellectuelle",
                "discussion": [
                    "Un point PI est à prévoir. Une preuve écrite indiquant que la DI est caduque est nécessaire auprès de la SATT.",
                    "Un rendez-vous avec Degive (INPI) et Kanpan (UMLP) est envisagé.",
                    "Une prestation Inno AER pour étude juridique PI pourrait être mobilisée, avec prise en charge à 50 %, plafonnée à 6 500 EUR HT.",
                    "Dans le cadre de l’incubation, des entretiens doivent être réalisés pour clarifier les clients potentiels : fournisseurs de matériel, fournisseurs de DPI, ou médecins prescripteurs.",
                ],
                "followup": "Programmer le point PI et préparer les entretiens d’incubation.",
            },
            {
                "subject": "Veille marché et brevets",
                "discussion": [
                    "La possibilité d’accéder aux bases de brevets via le laboratoire est à investiguer.",
                    "L’accès Xerfi via la BU et les rapports annuels de sociétés commercialisant des moniteurs sont évoqués.",
                    "Des rendez-vous business avec des fournisseurs pourraient aider à comprendre leurs solutions actuelles.",
                ],
                "followup": "Identifier les accès documentaires disponibles.",
            },
            {
                "subject": "Analyse CTG et variables",
                "discussion": [
                    "Gregory rappelle que le praticien extrait des features dans l’analyse CTG ; l’objectif serait d’automatiser une partie de cette extraction.",
                    "Les variables à extraire pourraient inclure : base, moyenne, minimum, maximum, nombre d’accélérations et nombre de décélérations.",
                    "La démarche est rapprochée de DeepCTG.",
                ],
                "followup": "",
            },
            {
                "subject": "Données synthétiques CARMENTA",
                "discussion": [
                    "Salma présente une génération de dataset synthétique à partir de CARMENTA, selon les règles métiers, distributions et types de variables.",
                    "Un jeu de 200 000 lignes est généré avec des NaN, dans les bornes du dictionnaire.",
                    "Les variables dupliquées sont conservées pour la traçabilité. Le document est composé de plusieurs tableaux et non d’un seul tableau.",
                    "Le nettoyage n’est pas encore réalisé ; les colonnes avec valeurs manquantes ou catégories peu claires sont générées avec NaN.",
                    "Gregory propose de supprimer les cas lorsque la donnée manque, car mettre 0 constituerait une information différente de l’absence d’information.",
                ],
                "followup": "Clarifier les ambiguïtés sémantiques du dictionnaire.",
            },
            {
                "subject": "Score clinico-biologique",
                "discussion": [
                    "L’issue défavorable est discutée à partir des variables biologiques : pH, base excess, lactates, delta pH artériel/veineux, PCO2 et PO2.",
                    "Les variables cliniques discutées sont les scores d’Apgar, le poids du nouveau-né et la température.",
                    "L’objectif du score CB est de produire un paramètre quantitatif du bien-être néonatal et de vérifier son association au risque de transfert.",
                    "Le décès et la réanimation définissent les issues les plus défavorables.",
                    "Le score doit combiner biologique et clinique, avec pondération en faveur de la clinique.",
                ],
                "followup": "Définir les seuils et classes retenus pour le score.",
            },
        ],
    },
    {
        "date": "26.01.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Suivi administratif",
                "discussion": [
                    "Une relance est nécessaire pour la signature du protocole, auprès de M. Kanpan ou Daussy.",
                    "Pour INTERREG, la FPP est en cours avec Jana Landry. La liste des membres du consortium, les noms, le nombre et le type de postes sont attendus.",
                    "Les disponibilités du bureau de Nicolas sont à clarifier.",
                ],
                "followup": "Relancer la signature du protocole et compléter les informations consortium.",
            },
            {
                "subject": "Questionnaire IA",
                "discussion": [
                    "Le questionnaire de familiarité avec l’IA est disponible via Google Forms.",
                    "Un rendez-vous PI avec M. Kanpan est à programmer, avec l’INPI en parallèle pour clore le sujet.",
                ],
                "followup": "Programmer le rendez-vous PI.",
            },
            {
                "subject": "Variables CARMENTA",
                "discussion": [
                    "Gregory interroge les variables pouvant être exclues par défaut : OBST, post, PMSI, BIO, fusion des parto S1/S2, doublons Echo, champs libres et variables Echo trop fines.",
                    "Une étude statistique sur le dictionnaire doit permettre d’identifier les paramètres ayant un effet et de classer les variables d’entrée.",
                    "Une colonne supplémentaire pourrait indiquer si la variable est conservée ou non.",
                ],
                "followup": "Ajouter une colonne de décision de conservation des variables.",
            },
            {
                "subject": "Équipe projet",
                "discussion": [
                    "L’équipe projet envisagée comprend Nicolas, Justine, une sage-femme recherche dédiée à 100 % ETP, un technicien de recherche clinique, deux stagiaires et un project manager DRCI.",
                ],
                "followup": "",
            },
        ],
    },
    {
        "date": "02.02.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Suivi Gregory",
                "discussion": [
                    "Le cahier des charges reste à clarifier. Le dernier courriel n’a pas été compris.",
                ],
                "followup": "Clarifier le cahier des charges avec Gregory.",
            },
            {
                "subject": "Administration et INTERREG",
                "discussion": [
                    "La relance concernant la signature de la convention a été faite le 27/01.",
                    "Pour le CHU, la liste des membres du consortium, les noms, le nombre et les types de postes sont attendus.",
                    "La sage-femme recherche et le technicien de recherche sont affectés au projet sur toute sa durée.",
                    "Les packs numériques ne sont pas nécessaires.",
                ],
                "followup": "Compléter les informations consortium CHU.",
            },
            {
                "subject": "Présentation Vouga",
                "discussion": [
                    "Les slides Vouga sont validées.",
                    "Une slide supplémentaire doit expliquer l’intérêt d’une nouvelle solution de surveillance du travail, dépassant les outils actuels et intégrant mieux le contexte obstétrical.",
                    "Le message clé est qu’un outil complémentaire produisant un score clinico-biologique est nécessaire face aux limites actuelles et à l’évolution des comorbidités obstétricales.",
                ],
                "followup": "Ajouter la slide de contexte et d’ambition.",
            },
            {
                "subject": "Questionnaire IA",
                "discussion": [
                    "Le questionnaire de familiarité avec l’IA est à articuler avec les sages-femmes recherche, le logiciel SPHINX et la procédure d’envoi.",
                    "Les questions doivent être envoyées à Nicolas avec copie à l’adresse sage-femme-recherche.",
                ],
                "followup": "Envoyer les questions à Nicolas et à l’adresse sage-femme-recherche.",
            },
            {
                "subject": "Données et issue défavorable",
                "discussion": [
                    "Les seuils biologiques et cliniques discutés le 12/01 sont repris.",
                    "NM indique qu’il faut identifier les variables pertinentes pour définir l’état du nouveau-né.",
                    "Les pires issues citées sont l’anoxo-ischémie, l’infection materno-foetale et la détresse respiratoire.",
                    "Les données de la base BIO doivent être utilisées pour pH et lactates, en supprimant les données de pH dans DIAMM.",
                ],
                "followup": "Utiliser BIO pour pH/lactates et retirer les pH DIAMM.",
            },
            {
                "subject": "Réunion DRCI/DPO",
                "discussion": [
                    "Pour le rendez-vous du 16/02 avec la DRCI et le DPO, il est proposé d’inviter M. Zerr (RSSI CHRUB).",
                ],
                "followup": "Inviter M. Zerr.",
            },
        ],
    },
    {
        "date": "09.02.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Suivi réunion",
                "discussion": [
                    "Le rendez-vous du 16/02 est déplacé.",
                ],
                "followup": "Reprogrammer le rendez-vous.",
            }
        ],
    },
    {
        "date": "23.02.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Point DSN et budget",
                "discussion": [
                    "Les contacts DSN à mettre dans la fiche de coût INTERREG sont Alexandra Monnot et Stéphane Tesolari.",
                    "Le budget d’un matériel pour l’exploitation locale et sa durée d’amortissement doivent être clarifiés avec l’UMLP.",
                    "Les questions portent aussi sur l’accès distant et le matériel informatique local.",
                ],
                "followup": "Valider le budget matériel avec l’UMLP.",
            },
            {
                "subject": "Convention CHU-UMLP",
                "discussion": [
                    "L’ajout de GG et SE à la convention CHU-UMLP est évoqué.",
                ],
                "followup": "Vérifier l’ajout de GG et SE.",
            },
            {
                "subject": "Architecture d’accès et calcul",
                "discussion": [
                    "Plusieurs options sont listées : accès distant sur la VM, apport de machine de calcul, mésocentre, machine UMLP, machine CHU ou machine via INTERREG.",
                    "L’accès distant CHU est validé.",
                ],
                "followup": "Arbitrer l’architecture technique cible.",
            },
            {
                "subject": "Consortium et INTERREG",
                "discussion": [
                    "NM doit refaire un mail à Manon Vouga pour valider le consortium.",
                    "Le CHUV doit contacter le canton de Vaud pour le cofinancement et Interreg CH pour avis sur le plan de financement actuel.",
                    "Une attention est portée au cadrage INTERREG : l’axe prioritaire porte sur la connaissance réciproque, les coopérations de formation, et l’écosystème interrégional d’innovation.",
                ],
                "followup": "NM envoie le mail à Manon Vouga.",
            },
        ],
    },
    {
        "date": "02.03.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Questionnaire sages-femmes",
                "discussion": [
                    "Justine prend en charge le questionnaire, avec un focus sur les médecins et sages-femmes de salle d’accouchement.",
                    "Les thèmes retenus sont le niveau de connaissance en IA, la compréhension du développement des outils, l’acceptabilité d’un outil d’aide, et les domaines d’aide attendus.",
                    "Le périmètre envisagé est BFC + Suisse.",
                ],
                "followup": "Contacter la DRCI pour vérifier le processus applicable.",
            },
            {
                "subject": "Hébergement de données",
                "discussion": [
                    "Un devis HDS externe doit être préparé.",
                ],
                "followup": "Préparer le devis HDS externe.",
            },
            {
                "subject": "INTERREG",
                "discussion": [
                    "La signature CHUV puis UMLP est attendue.",
                    "Une première réponse Interreg est annoncée sous dix jours.",
                    "Le travail sur le dossier complet se poursuit, notamment la partie financière suisse.",
                    "Une prise de contact CHUV avec l’EPFL est envisagée.",
                ],
                "followup": "Poursuivre le dossier complet et préciser la réunion DRCI/DRVI.",
            },
            {
                "subject": "Étude CARMENTA",
                "discussion": [
                    "La définition de la sortie nouveau-né de salle de travail doit prendre en compte le seuil supérieur ou égal à 37 SA.",
                    "Pour le document de Justine, les personnes à citer sont Dr Vouga et Noura Dridi ou Amir Hajjam selon les conditions d’appartenance au laboratoire.",
                    "Une prospection avec l’EPFL doit être indiquée en commentaire.",
                ],
                "followup": "Intégrer la mention de prospection EPFL.",
            },
        ],
    },
    {
        "date": "09.03.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Informations générales",
                "discussion": [
                    "Le document d’accès distant a été envoyé à R. Zerr. Une réponse est attendue, avec relance prévue le lendemain.",
                    "Les dernières réunions d’incubation ont lieu jeudi prochain, puis une semaine dédiée est prévue du 23 au 26 mars.",
                    "Les slides CARMENTA sont à relire.",
                ],
                "followup": "Relancer R. Zerr si nécessaire.",
            },
            {
                "subject": "Compte rendu du 19/03",
                "discussion": [
                    "La question est posée de savoir si NM peut prendre en charge le compte rendu du 19/03.",
                    "NM demande s’il existe un horaire programmé, sinon la réunion se tiendra en visioconférence.",
                    "Les slides du montage du projet doivent être mises à jour.",
                ],
                "followup": "Confirmer l’organisation du 19/03.",
            },
            {
                "subject": "Questionnaire Justine",
                "discussion": [
                    "Le questionnaire vise les cliniciens et reste généraliste.",
                    "Les objectifs sont d’évaluer les connaissances sur l’IA en obstétrique, les sentiments favorables ou défavorables, l’intégration dans des conditions de travail complexes, la crainte d’être remplacé, et les paramètres souhaités dans une solution IA.",
                ],
                "followup": "Finaliser les objectifs du questionnaire.",
            },
            {
                "subject": "Stage Sadia",
                "discussion": [
                    "Le sujet de stage doit être préparé et envoyé à F. Auber.",
                ],
                "followup": "Préparer et envoyer le sujet de stage.",
            },
        ],
    },
    {
        "date": "16.03.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Slides et réunions",
                "discussion": [
                    "La relecture des slides CR est validée.",
                    "Pour l’échange CHU-UMLP et la reprise de contact avec la SATT, l’équipe attend le retour INTERREG.",
                    "Pour la réunion du jeudi, l’horaire de NM n’est pas encore précisé. OB contacte Soraya Martin.",
                ],
                "followup": "Attendre le retour INTERREG avant CHU-UMLP et SATT.",
            },
            {
                "subject": "Accès et signatures",
                "discussion": [
                    "La question de la signature liée à la fusion mésocentre des deux côtés est posée.",
                    "NM a relancé M. Zerr.",
                ],
                "followup": "Suivre la réponse de M. Zerr.",
            },
            {
                "subject": "Process Sadia",
                "discussion": [
                    "Le process proposé comprend : reprendre la présentation CARMENTA, produire le flowchart du tri clinique, intégrer BIO dans DIAMM, mettre à jour la présentation avec les données actualisées, puis préparer la publication.",
                    "OB prépare le cahier des charges statistique.",
                ],
                "followup": "Formaliser le cahier des charges statistique.",
            },
            {
                "subject": "Questionnaire Justine",
                "discussion": [
                    "La démarche doit être demandée à Ingrid.",
                ],
                "followup": "Contacter Ingrid.",
            },
        ],
    },
    {
        "date": "23.03.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Suivis en cours",
                "discussion": [
                    "M. Zerr doit être relancé.",
                    "La convention avec FC Innov est en cours ; les relances doivent être surveillées.",
                    "Le challenge DECA est en cours.",
                ],
                "followup": "Relancer Zerr et suivre FC Innov.",
            },
            {
                "subject": "INTERREG",
                "discussion": [
                    "Le dossier complet est reçu, avec accès à Synergie.",
                    "Jana complète le début du dossier.",
                    "La DRVI centralise pour le CHU et l’UMLP.",
                ],
                "followup": "Poursuivre la saisie Synergie.",
            },
        ],
    },
    {
        "date": "30.03.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Plateforme Synergie",
                "discussion": [
                    "La relecture de la saisie de Jana est à réaliser.",
                    "Les WP doivent être confirmés, puis les affectations par WP doivent être précisées.",
                    "Le retour de Manon Vouga concernant l’EPFL est attendu.",
                ],
                "followup": "Relire Synergie et confirmer les WP.",
            },
            {
                "subject": "RSSI et accès",
                "discussion": [
                    "Un retour de M. Zerr est attendu. Les accès doivent être envoyés aujourd’hui ou demain.",
                    "Le document complet doit être préparé puis renvoyé à M. Zerr.",
                ],
                "followup": "Relancer si aucune nouvelle le lendemain après-midi.",
            },
            {
                "subject": "Stage Sadia",
                "discussion": [
                    "La fiche d’accès Sadia doit être signée par NM.",
                    "La convention de stage est en cours de préparation chez Cécile.",
                ],
                "followup": "Faire signer la fiche d’accès.",
            },
            {
                "subject": "Livrables INTERREG",
                "discussion": [
                    "Les livrables actuels incluent : schéma de structuration des données documenté, référentiel de variables et métadonnées, note méthodologique sur réplicabilité et transposabilité, prototypes de recherche codéveloppés, recommandations bilatérales et feuille de route scientifique.",
                ],
                "followup": "",
            },
            {
                "subject": "Travaux Selma",
                "discussion": [
                    "Selma travaille sur les données de synthèse : génération d’un dataset de 200 000 lignes et 441 variables, avec bornes et distributions du dictionnaire.",
                    "Après filtre >= 37 SA, 124 000 cas sont étudiés.",
                    "Une variable binaire de transfert en réanimation est créée pour préparer un modèle de prédiction, sans résultat concluant à ce stade.",
                    "Une génération à partir de 100 cas de données réelles est proposée.",
                    "NM signale des écarts de réalisme : 19 % de transfert incluant l’UK, variable sexe foetal à prendre après accouchement dans DIAMM, état ERCF à l’instant T du test d’admission, et répartition de prématurité non conforme.",
                ],
                "followup": "Corriger la définition du transfert et les variables sources.",
            },
        ],
    },
    {
        "date": "06.04.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Données ENP",
                "discussion": [
                    "La disponibilité des données ENP en CSV auprès de NM est à vérifier.",
                    "Si elles ne sont pas disponibles, l’équipe se demande si un site dynamique aiderait à produire l’analyse.",
                    "La question du niveau de détail attendu dans les tableaux est posée.",
                ],
                "followup": "Vérifier la disponibilité des CSV ENP.",
            },
            {
                "subject": "Sadia",
                "discussion": [
                    "La signature de l’habilitation Sadia est à suivre.",
                ],
                "followup": "Faire signer l’habilitation.",
            },
        ],
    },
    {
        "date": "20.04.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "HEIG-VD et gouvernance",
                "discussion": [
                    "La possibilité que HEIG-VD soit chef de file ou leader partenaire sur un WP est discutée.",
                    "Le processus de signature avec HEIG-VD doit être clarifié.",
                    "La question d’une nouvelle signature de la FPP est posée.",
                ],
                "followup": "Clarifier le processus de signature HEIG-VD.",
            },
            {
                "subject": "Prochaines étapes Stephan",
                "discussion": [
                    "L’annexe 4 doit être traitée, avec plafonds de financement : Interreg 146 537,60 EUR et Canton 90 000 EUR.",
                    "L’annexe 2 sur l’autofinancement sera préparée après validation de l’annexe 4.",
                    "L’annexe 6 concernant les références bancaires est également à produire.",
                ],
                "followup": "Finaliser l’annexe 4, puis préparer les annexes 2 et 6.",
            },
        ],
    },
    {
        "date": "27.04.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Travaux Gregory",
                "discussion": [
                    "Gregory adapte le format de C à un traitement algorithmique, avec intégration dans un vecteur.",
                    "Les clés primaires et secondaires sont définies : numéro d’accouchement en clé primaire, numéro permanent + date en clé secondaire.",
                    "Les paramètres d’entrée et de sortie sont triés.",
                    "Un post-traitement ERCF est prévu pour intégrer ces données dans C.",
                    "Les variables inutiles sont supprimées : valeurs 100 % identiques, numéro d’admission, etc.",
                    "La base de traitement est créée avec uniquement les entrées et sorties intéressantes. Après filtre 37 SA, environ 14 000 cas sont disponibles.",
                ],
                "followup": "Vérifier si la base pourra sortir du CHUB.",
            },
            {
                "subject": "Contrôle qualité ERCF",
                "discussion": [
                    "Un contrôle qualité doit vérifier que la durée de l’ERCF est cohérente avec la durée du travail.",
                ],
                "followup": "Mettre en place le contrôle de cohérence des durées.",
            },
            {
                "subject": "Score clinico-biologique",
                "discussion": [
                    "Le transfert en réanimation est actuellement considéré comme le juge de paix pour l’état défavorable.",
                    "Les obstétriciens sont sensibles à l’adaptation néonatale lors de l’accouchement, évaluée par Apgar et gazométrie au cordon.",
                    "La prédiction du pH seul est jugée insuffisante, car un mauvais pH peut coexister avec une bonne adaptation clinique.",
                    "Le couplage des variables clinico-biologiques doit aider la décision du praticien.",
                    "Pour chaque enfant de C, l’objectif est d’obtenir une valeur de score, puis de relier ce score aux issues.",
                ],
                "followup": "Définir la construction opérationnelle du score.",
            },
            {
                "subject": "Données BIO et doublons",
                "discussion": [
                    "Si plusieurs valeurs existent dans BIO, la valeur la plus proche de l’accouchement est retenue.",
                    "Les doublons NIP dans DIAMM sont à investiguer ; BB1 et BB2 pourraient correspondre à des jumeaux.",
                ],
                "followup": "Analyser les doublons NIP dans DIAMM.",
            },
        ],
    },
    {
        "date": "04.05.2026",
        "title": "INOVIA weekly",
        "topics": [
            {
                "subject": "Sadia",
                "discussion": [
                    "L’état d’avancement de Sadia et l’extraction de données depuis le CHU sont à vérifier.",
                    "Sadia repart des fichiers sources pour nettoyer et prétraiter CARMENTA.",
                    "Le pipeline comprend : filtrage > 37 SA et dates d’inclusion, suppression des doublons, traitement des valeurs manquantes, analyse des identifiants absents d’une base ou d’une autre, et listing des paramètres d’entrée/sortie candidats pour prédire les complications.",
                ],
                "followup": "Suivre l’extraction CHU et le pipeline CARMENTA.",
            },
            {
                "subject": "Salma",
                "discussion": [
                    "L’état d’avancement de Salma et l’accueil via le fonds de la bourse BFT Lab sont à vérifier.",
                    "OB indique que la convention d’accueil est en cours et devrait se clarifier dans la semaine.",
                ],
                "followup": "Suivre la convention d’accueil.",
            },
            {
                "subject": "Travail CTU UHB",
                "discussion": [
                    "Un travail avec CTU UHB est envisagé sur la prédiction avec UC et RC, en attendant l’accès aux données.",
                ],
                "followup": "",
            },
            {
                "subject": "Travaux Gregory",
                "discussion": [
                    "Les fichiers CTS en TXT sont réduits de 14 Go à 370 Mo.",
                    "Les colonnes de mesures RCF 1 sont conservées ; l’acquisition de 5 secondes est abandonnée au profit d’une moyenne des valeurs non nulles.",
                    "Le format reste compatible avec la mise en relation avec le partogramme, sans perte majeure attendue par rapport à 250 ms selon les travaux précédents.",
                    "La difficulté actuelle porte sur la mise en relation des données de partogramme avec l’enregistrement, du fait du volume de données.",
                    "NM précise que l’objectif n’est pas de donner la méthode du praticien, mais d’analyser ce qui ressort du travail de Gregory.",
                ],
                "followup": "Organiser un échange praticien sur la mise en relation partogramme/enregistrement.",
            },
            {
                "subject": "Prédiction et réduction de base",
                "discussion": [
                    "Les données pré-partum disponibles avant l’entrée en salle peuvent estimer un risque d’issue néonatale.",
                    "La base doit être mise dans un tableau 2D, puis les paramètres pertinents doivent être mis en relation avec les issues recherchées.",
                    "Si un paramètre n’a d’influence sur aucun paramètre de sortie, il pourra être supprimé afin de réduire la base.",
                    "Après modélisation des scores Apgar et biologiques, un score clinico-biologique avec valeurs seuils sera créé.",
                ],
                "followup": "Définir les étapes d’analyse après réorganisation de la base.",
            },
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width_pct(cell, pct_width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "pct")
    tc_w.set(qn("w:w"), str(pct_width))


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def set_table_autofit_to_window(table):
    table.autofit = True
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is not None:
        tbl_pr.remove(tbl_layout)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(cell, text, bold=False, color=None, size=8.5, align=None):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.text:
        p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)


def add_header(doc, meeting):
    day, month, year = meeting["date"].split(".")
    title = f"Compte rendu de la réunion hebdomadaire INOVIA du {day} {month} {year}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Présents : ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r = p.add_run("Non renseigné dans les notes sources.")
    r.font.name = "Arial"
    r.font.size = Pt(9)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Excusés : ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r = p.add_run("Non renseigné dans les notes sources.")
    r.font.name = "Arial"
    r.font.size = Pt(9)


def add_meeting_table(doc, meeting):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_autofit_to_window(table)

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    headers = ["Sujets", "Discussion - Décision", "Suivi\nÉchéance"]
    pct_widths = [850, 3400, 750]
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width_pct(cell, pct_widths[i])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        set_cell_margins(cell, 110, 100, 110, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_text(cell, label, bold=True, color=BLUE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    for topic in meeting["topics"]:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            set_cell_width_pct(cell, pct_widths[i])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if i == 0:
                set_cell_shading(cell, "FFFFFF")
            elif i == 2 and topic.get("followup"):
                set_cell_shading(cell, LIGHT_BLUE)

        add_text(row.cells[0], topic["subject"], bold=True, size=8.5)
        discussion_cell = row.cells[1]
        for idx, item in enumerate(topic["discussion"]):
            p = discussion_cell.paragraphs[0] if idx == 0 else discussion_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(item)
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
        add_text(row.cells[2], topic.get("followup") or "", size=8.2)


def build_docx():
    doc = Document()
    configure_doc(doc)
    for idx, meeting in enumerate(MEETINGS):
        if idx:
            doc.add_page_break()
        add_header(doc, meeting)
        add_meeting_table(doc, meeting)
    doc.save(DOCX_OUT)


def build_individual_docx():
    INDIVIDUAL_DIR.mkdir(parents=True, exist_ok=True)
    for meeting in MEETINGS:
        doc = Document()
        configure_doc(doc)
        add_header(doc, meeting)
        add_meeting_table(doc, meeting)
        day, month, year = meeting["date"].split(".")
        stem = f"{year[-2:]}{month}{day}"
        doc.save(INDIVIDUAL_DIR / f"{stem}_CR_INOVIA_weekly.docx")


def build_md():
    lines = ["# Comptes rendus INOVIA hebdomadaires 2026", ""]
    for meeting in MEETINGS:
        lines += [
            f"## PV du {meeting['date']} - {meeting['title']}",
            "",
            "**Présents :** Non renseigné dans les notes sources.",
            "",
            "**Excusés :** Non renseigné dans les notes sources.",
            "",
            "| Sujets | Discussion - Décision | Suivi / Échéance |",
            "| --- | --- | --- |",
        ]
        for topic in meeting["topics"]:
            discussion = "<br>".join(topic["discussion"])
            followup = topic.get("followup", "")
            lines.append(f"| {topic['subject']} | {discussion} | {followup} |")
        lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_md()
    build_docx()
    build_individual_docx()
    print(DOCX_OUT)
    print(MD_OUT)
    print(INDIVIDUAL_DIR)


if __name__ == "__main__":
    main()
